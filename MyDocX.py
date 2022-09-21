from docx import Document
from docx.shared import Inches
import docx

class MyDocX:
    def __init__(self, fileName, title) -> None:
        
        self.FileName = "Data/" + fileName + ".docx"
        self.Title = title
        self.Document = Document()
        self.Document.add_heading(f'{title}', 0)
        doc = docx.Document()
        doc.save(f'{self.FileName}')
    
    def AddContent(self, content):
        try:
            doc = docx.Document(f'{self.FileName}')
            p = doc.add_paragraph(f'{content}')
        except:
            docx.save(f'{self.FileName}')
        # p.add_run('bold').bold = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True

        # document.add_heading('Heading, level 1', level=1)
        # document.add_paragraph('Intense quote', style='Intense Quote')

        # document.add_paragraph(
        #     'first item in unordered list', style='List Bullet'
        # )
        # document.add_paragraph(
        #     'first item in ordered list', style='List Number'
        # )

        # document.add_picture('clock.png', width=Inches(1.25))

        # records = (
        #     (3, '101', 'Spam'),
        #     (7, '422', 'Eggs'),
        #     (4, '631', 'Spam, spam, eggs, and spam')
        # )

        # table = document.add_table(rows=1, cols=3)
        # hdr_cells = table.rows[0].cells
        # hdr_cells[0].text = 'Qty'
        # hdr_cells[1].text = 'Id'
        # hdr_cells[2].text = 'Desc'
        # for qty, id, desc in records:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = str(qty)
        #     row_cells[1].text = id
        #     row_cells[2].text = desc
        doc.add_page_break()
        doc.save(f'{self.FileName}')
    
    
# doc = DocX("Moto", "Test")
# for i in range(0,2):
#     doc.AddContent(f'Thêm lần thứ {i}')